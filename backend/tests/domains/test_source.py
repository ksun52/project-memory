import io
import uuid

import pytest

from app.domains.auth.models import User
from app.domains.memory_space.models import MemorySpace
from app.domains.source.models import (
    Source,
    SourceChunk,
    SourceContent,
    SourceContentEntity,
    SourceCreateDocument,
    SourceCreateNote,
    SourceEntity,
    SourceFile,
    SourceFileEntity,
    SourceChunkEntity,
)
from app.domains.workspace.models import Workspace

DEV_USER_ID = uuid.UUID("00000000-0000-0000-0000-000000000001")
OTHER_USER_ID = uuid.UUID("00000000-0000-0000-0000-000000000002")


def _seed_dev_user(db_session) -> None:
    user = User(
        id=DEV_USER_ID,
        auth_provider="dev",
        auth_provider_id="dev-user-001",
        email="dev@projectmemory.local",
        display_name="Dev User",
    )
    db_session.add(user)
    db_session.commit()


def _seed_other_user(db_session) -> None:
    user = User(
        id=OTHER_USER_ID,
        auth_provider="dev",
        auth_provider_id="dev-user-002",
        email="other@projectmemory.local",
        display_name="Other User",
    )
    db_session.add(user)
    db_session.commit()


def _create_workspace_and_space(db_session, owner_id=DEV_USER_ID):
    workspace = Workspace(
        owner_id=owner_id,
        name="Test Workspace",
        description="",
    )
    db_session.add(workspace)
    db_session.commit()
    db_session.refresh(workspace)

    ms = MemorySpace(
        workspace_id=workspace.id,
        name="Test Space",
        description="",
        status="active",
    )
    db_session.add(ms)
    db_session.commit()
    db_session.refresh(ms)
    return workspace, ms


# ===== Entity Tests =====


class TestSourceEntity:
    def test_from_orm(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        source = Source(
            memory_space_id=ms.id,
            source_type="note",
            title="Test Note",
            processing_status="pending",
        )
        db_session.add(source)
        db_session.commit()
        db_session.refresh(source)

        entity = SourceEntity.from_orm(source)
        assert entity.id == source.id
        assert entity.memory_space_id == ms.id
        assert entity.source_type == "note"
        assert entity.title == "Test Note"
        assert entity.processing_status == "pending"
        assert entity.processing_error is None
        assert entity.created_at is not None
        assert entity.updated_at is not None


class TestSourceContentEntity:
    def test_fields(self):
        entity = SourceContentEntity(
            source_id=uuid.uuid4(),
            content_text="Hello world",
        )
        assert entity.content_text == "Hello world"


class TestSourceFileEntity:
    def test_fields(self):
        entity = SourceFileEntity(
            source_id=uuid.uuid4(),
            mime_type="application/pdf",
            size_bytes=1024,
            original_filename="test.pdf",
        )
        assert entity.mime_type == "application/pdf"
        assert entity.size_bytes == 1024


class TestSourceChunkEntity:
    def test_fields(self):
        entity = SourceChunkEntity(
            id=uuid.uuid4(),
            source_id=uuid.uuid4(),
            chunk_index=0,
            content="chunk text",
            start_offset=0,
            end_offset=10,
        )
        assert entity.chunk_index == 0


# ===== Schema Validation Tests =====


class TestSourceCreateNote:
    def test_valid(self):
        schema = SourceCreateNote(source_type="note", title="My Note", content="Some content")
        assert schema.title == "My Note"
        assert schema.content == "Some content"

    def test_empty_title_rejected(self):
        with pytest.raises(Exception):
            SourceCreateNote(source_type="note", title="   ", content="content")

    def test_empty_content_rejected(self):
        with pytest.raises(Exception):
            SourceCreateNote(source_type="note", title="title", content="   ")

    def test_wrong_source_type_rejected(self):
        with pytest.raises(Exception):
            SourceCreateNote(source_type="document", title="title", content="content")


class TestSourceCreateDocument:
    def test_valid(self):
        schema = SourceCreateDocument(source_type="document", title="My Doc")
        assert schema.title == "My Doc"

    def test_empty_title_rejected(self):
        with pytest.raises(Exception):
            SourceCreateDocument(source_type="document", title="   ")


# ===== Service Tests =====


class TestCreateNoteSource:
    def test_creates_source_and_content(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="My Note", content="Note body text")
        entity = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        assert isinstance(entity, SourceEntity)
        assert entity.source_type == "note"
        assert entity.title == "My Note"
        assert entity.processing_status == "pending"

        # Verify SourceContent was created
        sc = db_session.query(SourceContent).filter(
            SourceContent.source_id == entity.id
        ).first()
        assert sc is not None
        assert sc.content_text == "Note body text"

    def test_rejects_nonexistent_memory_space(self, db_session):
        _seed_dev_user(db_session)
        from app.domains.source import service
        from app.core.exceptions import NotFoundError

        data = SourceCreateNote(source_type="note", title="Note", content="body")
        with pytest.raises(NotFoundError):
            service.create_note_source(db_session, uuid.uuid4(), DEV_USER_ID, data)

    def test_rejects_other_user(self, db_session):
        _seed_dev_user(db_session)
        _seed_other_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service
        from app.core.exceptions import ForbiddenError

        data = SourceCreateNote(source_type="note", title="Note", content="body")
        with pytest.raises(ForbiddenError):
            service.create_note_source(db_session, ms.id, OTHER_USER_ID, data)


class TestCreateDocumentSource:
    def test_creates_source_and_file(self, db_session, tmp_path):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service
        from unittest.mock import MagicMock, patch

        # Create a fake uploaded file (plain text)
        file_content = b"Hello, this is a plain text document."
        fake_file = MagicMock()
        fake_file.filename = "test.txt"
        fake_file.content_type = "text/plain"
        fake_file.file = io.BytesIO(file_content)
        fake_file.size = len(file_content)

        with patch.object(service, "storage_client") as mock_storage, \
             patch.object(service, "parse_document") as mock_parse:
            mock_storage.save_file.return_value = "/storage/path/test.txt"
            mock_parse.return_value = "Hello, this is a plain text document."

            entity = service.create_document_source(
                db_session, ms.id, DEV_USER_ID, "Test Doc", fake_file
            )

        assert isinstance(entity, SourceEntity)
        assert entity.source_type == "document"
        assert entity.title == "Test Doc"
        assert entity.processing_status == "pending"

        # Verify SourceFile was created
        sf = db_session.query(SourceFile).filter(
            SourceFile.source_id == entity.id
        ).first()
        assert sf is not None
        assert sf.mime_type == "text/plain"
        assert sf.original_filename == "test.txt"

        # Verify SourceContent was created from parsed text
        sc = db_session.query(SourceContent).filter(
            SourceContent.source_id == entity.id
        ).first()
        assert sc is not None
        assert "Hello" in sc.content_text


class TestParseDocument:
    def test_parse_txt(self, tmp_path):
        from app.domains.source.service import parse_document

        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Hello world")
        result = parse_document(str(txt_file), "text/plain")
        assert result == "Hello world"

    def test_parse_unsupported_type(self, tmp_path):
        from app.domains.source.service import parse_document
        from app.core.exceptions import ValidationError

        html_file = tmp_path / "test.html"
        html_file.write_text("<html>hi</html>")
        with pytest.raises(ValidationError):
            parse_document(str(html_file), "text/html")

    def test_parse_pdf(self, tmp_path):
        """Test PDF parsing with a real simple PDF."""
        from app.domains.source.service import parse_document

        # Create a minimal PDF using reportlab if available, otherwise skip
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed, skipping PDF test")

        pdf_path = tmp_path / "test.pdf"
        c = canvas.Canvas(str(pdf_path))
        c.drawString(72, 720, "Hello PDF World")
        c.save()

        result = parse_document(str(pdf_path), "application/pdf")
        assert "Hello PDF World" in result

    def test_parse_docx(self, tmp_path):
        """Test DOCX parsing."""
        from app.domains.source.service import parse_document
        import docx

        docx_path = tmp_path / "test.docx"
        doc = docx.Document()
        doc.add_paragraph("Hello DOCX World")
        doc.save(str(docx_path))

        result = parse_document(str(docx_path), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert "Hello DOCX World" in result


class TestChunkSourceContent:
    def test_single_chunk_for_short_content(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        source = Source(
            memory_space_id=ms.id,
            source_type="note",
            title="Short",
            processing_status="pending",
        )
        db_session.add(source)
        db_session.commit()
        db_session.refresh(source)

        from app.domains.source.service import chunk_source_content

        chunks = chunk_source_content(db_session, source.id, "Short content.")
        assert len(chunks) == 1
        assert chunks[0].content == "Short content."
        assert chunks[0].start_offset == 0
        assert chunks[0].end_offset == len("Short content.")
        assert chunks[0].chunk_index == 0

    def test_multiple_chunks_with_overlap(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        source = Source(
            memory_space_id=ms.id,
            source_type="note",
            title="Long",
            processing_status="pending",
        )
        db_session.add(source)
        db_session.commit()
        db_session.refresh(source)

        from app.domains.source.service import chunk_source_content

        # Create long content that exceeds chunk target
        sentences = [f"This is sentence number {i}. " for i in range(200)]
        long_text = "".join(sentences)

        chunks = chunk_source_content(db_session, source.id, long_text)
        assert len(chunks) > 1

        # Verify chunk indices are sequential
        for i, chunk in enumerate(chunks):
            assert chunk.chunk_index == i

        # Verify offsets are correct (start of first = 0, end of last = len)
        assert chunks[0].start_offset == 0
        assert chunks[-1].end_offset == len(long_text)

        # Verify chunks are stored in DB
        db_chunks = db_session.query(SourceChunk).filter(
            SourceChunk.source_id == source.id
        ).order_by(SourceChunk.chunk_index).all()
        assert len(db_chunks) == len(chunks)

    def test_overlap_exists_between_chunks(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        source = Source(
            memory_space_id=ms.id,
            source_type="note",
            title="Overlap Test",
            processing_status="pending",
        )
        db_session.add(source)
        db_session.commit()
        db_session.refresh(source)

        from app.domains.source.service import chunk_source_content

        sentences = [f"Sentence {i} of the test document. " for i in range(200)]
        long_text = "".join(sentences)

        chunks = chunk_source_content(db_session, source.id, long_text)

        if len(chunks) >= 2:
            # The second chunk should start before the first chunk ends
            # (overlap means start_offset of chunk N+1 < end_offset of chunk N)
            for i in range(len(chunks) - 1):
                # Overlap: next chunk starts before current ends in the original text
                assert chunks[i + 1].start_offset < chunks[i].end_offset, (
                    f"Chunk {i+1} start ({chunks[i+1].start_offset}) should be < "
                    f"chunk {i} end ({chunks[i].end_offset}) for overlap"
                )


class TestListSources:
    def test_list_empty(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        entities, total = service.list_sources(db_session, ms.id, DEV_USER_ID)
        assert entities == []
        assert total == 0

    def test_list_with_sources(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        for i in range(3):
            data = SourceCreateNote(source_type="note", title=f"Note {i}", content=f"Content {i}")
            service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        entities, total = service.list_sources(db_session, ms.id, DEV_USER_ID)
        assert total == 3
        assert len(entities) == 3

    def test_list_pagination(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        for i in range(5):
            data = SourceCreateNote(source_type="note", title=f"Note {i}", content=f"Content {i}")
            service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        entities, total = service.list_sources(db_session, ms.id, DEV_USER_ID, page=1, page_size=2)
        assert total == 5
        assert len(entities) == 2

    def test_list_filter_by_type(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        # Create notes
        for i in range(3):
            data = SourceCreateNote(source_type="note", title=f"Note {i}", content=f"Content {i}")
            service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        # Create a document source directly via ORM for filtering test
        doc_source = Source(
            memory_space_id=ms.id,
            source_type="document",
            title="Doc",
            processing_status="pending",
        )
        db_session.add(doc_source)
        db_session.commit()

        entities, total = service.list_sources(
            db_session, ms.id, DEV_USER_ID, source_type="note"
        )
        assert total == 3

    def test_list_filter_by_status(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="Note", content="Content")
        service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        entities, total = service.list_sources(
            db_session, ms.id, DEV_USER_ID, processing_status="pending"
        )
        assert total == 1

        entities, total = service.list_sources(
            db_session, ms.id, DEV_USER_ID, processing_status="completed"
        )
        assert total == 0


class TestGetSource:
    def test_get_source(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="My Note", content="Body")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        entity = service.get_source(db_session, created.id, DEV_USER_ID)
        assert entity.id == created.id
        assert entity.title == "My Note"

    def test_get_source_not_found(self, db_session):
        _seed_dev_user(db_session)
        from app.domains.source import service
        from app.core.exceptions import NotFoundError

        with pytest.raises(NotFoundError):
            service.get_source(db_session, uuid.uuid4(), DEV_USER_ID)

    def test_get_source_wrong_owner(self, db_session):
        _seed_dev_user(db_session)
        _seed_other_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service
        from app.core.exceptions import ForbiddenError

        data = SourceCreateNote(source_type="note", title="My Note", content="Body")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        with pytest.raises(ForbiddenError):
            service.get_source(db_session, created.id, OTHER_USER_ID)


class TestGetSourceDetail:
    def test_get_detail_with_content(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="Detailed Note", content="Detail body")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        detail = service.get_source_detail(db_session, created.id, DEV_USER_ID)
        assert detail["source"].id == created.id
        assert detail["content"].content_text == "Detail body"
        assert detail["file"] is None


class TestGetSourceContent:
    def test_get_content(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="Note", content="The content")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        content = service.get_source_content(db_session, created.id, DEV_USER_ID)
        assert content.content_text == "The content"


class TestDeleteSource:
    def test_delete_source(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service
        from app.core.exceptions import NotFoundError

        data = SourceCreateNote(source_type="note", title="Delete Me", content="Body")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        service.delete_source(db_session, created.id, DEV_USER_ID)

        # Should not be findable anymore
        with pytest.raises(NotFoundError):
            service.get_source(db_session, created.id, DEV_USER_ID)

    def test_delete_cascades_content(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="Delete Me", content="Body")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        service.delete_source(db_session, created.id, DEV_USER_ID)

        db_session.expire_all()
        sc = db_session.query(SourceContent).filter(
            SourceContent.source_id == created.id,
            SourceContent.deleted_at.is_(None),
        ).first()
        assert sc is None

    def test_delete_cascades_chunks(self, db_session):
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)

        from app.domains.source import service

        data = SourceCreateNote(source_type="note", title="Chunked", content="Body text")
        created = service.create_note_source(db_session, ms.id, DEV_USER_ID, data)

        # Create chunks
        service.chunk_source_content(db_session, created.id, "Body text")

        service.delete_source(db_session, created.id, DEV_USER_ID)

        db_session.expire_all()
        chunks = db_session.query(SourceChunk).filter(
            SourceChunk.source_id == created.id,
            SourceChunk.deleted_at.is_(None),
        ).all()
        assert len(chunks) == 0

    def test_delete_not_found(self, db_session):
        _seed_dev_user(db_session)
        from app.domains.source import service
        from app.core.exceptions import NotFoundError

        with pytest.raises(NotFoundError):
            service.delete_source(db_session, uuid.uuid4(), DEV_USER_ID)


# ===== Router Tests =====


class TestSourceRouter:
    def _setup(self, client, db_session):
        """Seed user, workspace, and memory space for router tests."""
        _seed_dev_user(db_session)
        workspace, ms = _create_workspace_and_space(db_session)
        return ms

    def test_create_note_via_api(self, client, db_session):
        ms = self._setup(client, db_session)
        resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={
                "source_type": "note",
                "title": "API Note",
                "content": "API note content",
            },
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["source_type"] == "note"
        assert data["title"] == "API Note"
        assert data["processing_status"] == "pending"
        assert "id" in data

    def test_create_document_via_api(self, client, db_session):
        ms = self._setup(client, db_session)
        file_content = b"This is a plain text document for testing."
        resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            data={"title": "API Doc"},
            files={"file": ("test.txt", io.BytesIO(file_content), "text/plain")},
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["source_type"] == "document"
        assert data["title"] == "API Doc"

    def test_list_sources_via_api(self, client, db_session):
        ms = self._setup(client, db_session)

        # Create a couple sources
        for i in range(3):
            client.post(
                f"/api/v1/memory-spaces/{ms.id}/sources",
                json={
                    "source_type": "note",
                    "title": f"Note {i}",
                    "content": f"Content {i}",
                },
            )

        resp = client.get(f"/api/v1/memory-spaces/{ms.id}/sources")
        assert resp.status_code == 200
        data = resp.json()
        assert data["total"] == 3
        assert len(data["items"]) == 3

    def test_list_sources_pagination(self, client, db_session):
        ms = self._setup(client, db_session)
        for i in range(5):
            client.post(
                f"/api/v1/memory-spaces/{ms.id}/sources",
                json={"source_type": "note", "title": f"Note {i}", "content": f"C {i}"},
            )

        resp = client.get(f"/api/v1/memory-spaces/{ms.id}/sources", params={"page_size": 2})
        assert resp.status_code == 200
        data = resp.json()
        assert data["total"] == 5
        assert len(data["items"]) == 2

    def test_list_sources_filter_type(self, client, db_session):
        ms = self._setup(client, db_session)
        client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={"source_type": "note", "title": "Note", "content": "body"},
        )

        resp = client.get(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            params={"source_type": "document"},
        )
        assert resp.status_code == 200
        assert resp.json()["total"] == 0

    def test_get_source_detail_via_api(self, client, db_session):
        ms = self._setup(client, db_session)
        create_resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={"source_type": "note", "title": "Detail Note", "content": "content"},
        )
        source_id = create_resp.json()["id"]

        resp = client.get(f"/api/v1/sources/{source_id}")
        assert resp.status_code == 200
        data = resp.json()
        assert data["id"] == source_id
        assert data["content"]["content_text"] == "content"
        assert data["file"] is None

    def test_get_source_content_via_api(self, client, db_session):
        ms = self._setup(client, db_session)
        create_resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={"source_type": "note", "title": "Content Note", "content": "my content"},
        )
        source_id = create_resp.json()["id"]

        resp = client.get(f"/api/v1/sources/{source_id}/content")
        assert resp.status_code == 200
        data = resp.json()
        assert data["content_text"] == "my content"

    def test_delete_source_via_api(self, client, db_session):
        ms = self._setup(client, db_session)
        create_resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={"source_type": "note", "title": "Delete Me", "content": "body"},
        )
        source_id = create_resp.json()["id"]

        resp = client.delete(f"/api/v1/sources/{source_id}")
        assert resp.status_code == 204

        # Verify it's gone
        resp = client.get(f"/api/v1/sources/{source_id}")
        assert resp.status_code == 404

    def test_get_source_not_found(self, client, db_session):
        _seed_dev_user(db_session)
        resp = client.get(f"/api/v1/sources/{uuid.uuid4()}")
        assert resp.status_code == 404

    def test_create_note_missing_content(self, client, db_session):
        ms = self._setup(client, db_session)
        resp = client.post(
            f"/api/v1/memory-spaces/{ms.id}/sources",
            json={"source_type": "note", "title": "No Content"},
        )
        assert resp.status_code == 422
